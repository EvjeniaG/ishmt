from __future__ import annotations

import json
import sys
from pathlib import Path

from docx import Document
from openpyxl import load_workbook
from pypdf import PdfReader


def docx_summary(path: Path) -> dict:
    doc = Document(path)
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    tables = []
    for table in doc.tables:
        rows = []
        for row in table.rows:
            cells = [cell.text.strip().replace("\n", " | ") for cell in row.cells]
            if any(cells):
                rows.append(cells)
        if rows:
            tables.append(rows)
    return {
        "file": str(path),
        "type": "docx",
        "paragraphs": paragraphs,
        "tables": tables,
    }


def pdf_summary(path: Path, max_pages: int = 80) -> dict:
    reader = PdfReader(str(path))
    pages = []
    for index, page in enumerate(reader.pages[:max_pages], start=1):
        text = (page.extract_text() or "").strip()
        if text:
            pages.append({"page": index, "text": text})
    return {
        "file": str(path),
        "type": "pdf",
        "page_count": len(reader.pages),
        "pages": pages,
    }


def xlsx_summary(path: Path) -> dict:
    workbook = load_workbook(path, data_only=False)
    sheets = []
    for sheet in workbook.worksheets:
        rows = []
        for row in sheet.iter_rows():
            values = []
            for cell in row:
                value = cell.value
                if value is None:
                    values.append("")
                else:
                    values.append(str(value))
            if any(values):
                rows.append(values)
        sheets.append(
            {
                "name": sheet.title,
                "max_row": sheet.max_row,
                "max_column": sheet.max_column,
                "rows": rows[:200],
            }
        )
    return {
        "file": str(path),
        "type": "xlsx",
        "sheets": sheets,
    }


def main() -> None:
    outputs = []
    for raw in sys.argv[1:]:
        path = Path(raw)
        suffix = path.suffix.lower()
        if suffix == ".docx":
            outputs.append(docx_summary(path))
        elif suffix == ".pdf":
            outputs.append(pdf_summary(path))
        elif suffix == ".xlsx":
            outputs.append(xlsx_summary(path))
        else:
            outputs.append({"file": str(path), "error": f"Unsupported suffix {suffix}"})
    print(json.dumps(outputs, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
